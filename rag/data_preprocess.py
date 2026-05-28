from __future__ import annotations

"""RAG 数据预处理模块。

本文件只负责 RAG 入库前的数据预处理能力：
- 获取文档文本：校验并读取 `.docx` 文档中的段落和表格文本
- 清洗文档文本：去除多余空白，保留必要段落边界
- 文本分块：使用中文友好的递归字符切分策略生成 chunks
- 文本向量化：调用 DashScope embedding 模型生成 chunk 向量
- 向量存储：将分块文本、向量和元数据持久化到 Chroma 数据库

说明：
- `python-docx` 不能直接读取旧版 `.doc` 二进制文档，遇到 `.doc` 会提示先转换为 `.docx`
- 需要提前在 `.env` 中配置 `DASHSCOPE_API_KEY`
"""

import os
from dataclasses import dataclass
from pathlib import Path
from typing import Sequence

import chromadb
from chromadb import Collection
from dotenv import load_dotenv
from docx import Document as DocxDocument
from langchain_community.embeddings import DashScopeEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

SUPPORTED_SUFFIXES = {".doc", ".docx"}
DEFAULT_CHUNK_SIZE = 500
DEFAULT_CHUNK_OVERLAP = 80
DEFAULT_EMBEDDING_MODEL = "text-embedding-v4"
DEFAULT_CHROMA_PERSIST_DIR = "./chroma_db"
DEFAULT_CHROMA_COLLECTION = "customer_service_docs"

# 面向中文文档设计的分隔符优先级：
# 1. 先按段落、换行等文档结构切分
# 2. 再按中文句末标点和分句标点切分
# 3. 兼容英文标点、空格
# 4. 最后退化为逐字符切分，避免超长文本无法切开
CHINESE_SEPARATORS: list[str] = [
    "\n\n",
    "\n",
    "。",
    "！",
    "？",
    "；",
    ";",
    "……",
    "…",
    "，",
    ",",
    "、",
    "：",
    ":",
    "）",
    ")",
    " ",
    "",
]


@dataclass(slots=True)
class TextChunk:
    """清洗并切分后的文本块。"""

    chunk_id: int
    text: str
    metadata: dict[str, str | int]


@dataclass(slots=True)
class VectorizedChunk:
    """已完成向量化的文本块。"""

    chunk_id: int
    text: str
    vector: list[float]
    metadata: dict[str, str | int | float | None]


def validate_document_path(file_path: str | Path) -> Path:
    """校验输入文件路径和文档类型。"""
    path = Path(file_path).expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")
    if not path.is_file():
        raise ValueError(f"输入路径不是文件: {path}")
    if path.suffix.lower() not in SUPPORTED_SUFFIXES:
        raise ValueError(f"仅支持 doc/docx 格式文档: {path}")
    return path


def normalize_text(text: str) -> str:
    """清洗文档文本，减少无意义空行并保留段落边界。"""
    lines = [line.strip() for line in text.splitlines()]
    normalized_lines: list[str] = []
    previous_blank = False

    for line in lines:
        if not line:
            if not previous_blank and normalized_lines:
                normalized_lines.append("")
            previous_blank = True
            continue

        normalized_lines.append(line)
        previous_blank = False

    return "\n".join(normalized_lines).strip()


def read_docx_text(file_path: str | Path) -> str:
    """使用 python-docx 读取 docx 段落和表格文本。"""
    path = validate_document_path(file_path)
    if path.suffix.lower() == ".doc":
        raise ValueError("python-docx 不支持 .doc 文件，请先转换为 .docx")

    document = DocxDocument(str(path))
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append("\t".join(cells))

    return normalize_text("\n".join(text_parts))


def split_text(
    text: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
    separators: Sequence[str] | None = None,
) -> list[str]:
    """使用 RecursiveCharacterTextSplitter 对中文文本进行递归字符切分。"""
    cleaned = normalize_text(text)
    if not cleaned:
        return []
    if chunk_size <= 0:
        raise ValueError("chunk_size 必须大于 0")
    if chunk_overlap < 0:
        raise ValueError("chunk_overlap 不能为负数")
    if chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap 必须小于 chunk_size")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=list(separators or CHINESE_SEPARATORS),
        keep_separator=True,
        length_function=len,
        is_separator_regex=False,
    )
    return [chunk.strip() for chunk in splitter.split_text(cleaned) if chunk.strip()]


def build_text_chunks(chunks: Sequence[str], file_path: str | Path) -> list[TextChunk]:
    """为文本分块补充基础元数据。"""
    path = Path(file_path).expanduser().resolve()
    return [
        TextChunk(
            chunk_id=index,
            text=chunk,
            metadata={
                "file_name": path.name,
                "file_path": str(path),
                "chunk_id": index,
                "chunk_length": len(chunk),
            },
        )
        for index, chunk in enumerate(chunks)
    ]


def build_embeddings(model_name: str = DEFAULT_EMBEDDING_MODEL) -> DashScopeEmbeddings:
    """构建阿里通义千问向量模型实例。"""
    load_dotenv()
    api_key = os.getenv("DASHSCOPE_API_KEY")
    if not api_key or api_key == "请在这里填写你的阿里云百炼 / 通义千问 API Key":
        raise EnvironmentError("请先在 .env 中配置 DASHSCOPE_API_KEY")

    return DashScopeEmbeddings(
        model=model_name,
        dashscope_api_key=api_key,
    )


def vectorize_texts(
    texts: Sequence[str],
    embeddings: DashScopeEmbeddings | None = None,
    model_name: str = DEFAULT_EMBEDDING_MODEL,
) -> list[list[float]]:
    """将文本列表转换为 embedding 向量列表。"""
    embedding_model = embeddings or build_embeddings(model_name)
    vectors = embedding_model.embed_documents(list(texts))
    return [[float(value) for value in vector] for vector in vectors]


def vectorize_chunks(
    chunks: Sequence[TextChunk],
    embeddings: DashScopeEmbeddings | None = None,
    model_name: str = DEFAULT_EMBEDDING_MODEL,
) -> list[VectorizedChunk]:
    """将文本块转换为带向量和元数据的文本块。"""
    embedding_model = embeddings or build_embeddings(model_name)
    texts = [chunk.text for chunk in chunks]
    vectors = vectorize_texts(texts, embeddings=embedding_model)

    vectorized_chunks: list[VectorizedChunk] = []
    for chunk, vector in zip(chunks, vectors, strict=True):
        vector_norm = float(sum(value * value for value in vector) ** 0.5)
        vectorized_chunks.append(
            VectorizedChunk(
                chunk_id=chunk.chunk_id,
                text=chunk.text,
                vector=vector,
                metadata={
                    **chunk.metadata,
                    "embedding_model": embedding_model.model,
                    "embedding_dimension": len(vector),
                    "embedding_norm": vector_norm,
                },
            )
        )

    return vectorized_chunks


def build_chroma_collection(
    persist_directory: str | Path = DEFAULT_CHROMA_PERSIST_DIR,
    collection_name: str = DEFAULT_CHROMA_COLLECTION,
) -> Collection:
    """创建或获取 Chroma 持久化 collection。"""
    persist_path = Path(persist_directory).expanduser().resolve()
    persist_path.mkdir(parents=True, exist_ok=True)
    client = chromadb.PersistentClient(path=str(persist_path))
    return client.get_or_create_collection(
        name=collection_name,
        metadata={"hnsw:space": "cosine"},
    )


def build_chroma_ids(chunks: Sequence[VectorizedChunk]) -> list[str]:
    """根据文件路径和 chunk_id 生成稳定的 Chroma 文档 ID。"""
    ids: list[str] = []
    for chunk in chunks:
        file_path = str(chunk.metadata.get("file_path", "unknown_file"))
        ids.append(f"{file_path}::chunk_{chunk.chunk_id}")
    return ids


def store_chunks_to_chroma(
    chunks: Sequence[VectorizedChunk],
    persist_directory: str | Path = DEFAULT_CHROMA_PERSIST_DIR,
    collection_name: str = DEFAULT_CHROMA_COLLECTION,
    clear_existing: bool = True,
) -> Collection:
    """将已向量化的文本块写入 Chroma 数据库。

    默认会先删除同一源文件已存在的分块，避免重新预处理后旧 chunk 残留。
    """
    collection = build_chroma_collection(
        persist_directory=persist_directory,
        collection_name=collection_name,
    )
    if not chunks:
        return collection

    if clear_existing:
        file_paths = {
            str(chunk.metadata["file_path"])
            for chunk in chunks
            if "file_path" in chunk.metadata
        }
        for file_path in file_paths:
            collection.delete(where={"file_path": file_path})

    collection.upsert(
        ids=build_chroma_ids(chunks),
        documents=[chunk.text for chunk in chunks],
        embeddings=[chunk.vector for chunk in chunks],
        metadatas=[chunk.metadata for chunk in chunks],
    )
    return collection


def preprocess_document(
    file_path: str | Path,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
    separators: Sequence[str] | None = None,
    embeddings: DashScopeEmbeddings | None = None,
    model_name: str = DEFAULT_EMBEDDING_MODEL,
    persist_directory: str | Path = DEFAULT_CHROMA_PERSIST_DIR,
    collection_name: str = DEFAULT_CHROMA_COLLECTION,
    store_to_chroma: bool = True,
    clear_existing: bool = True,
) -> list[VectorizedChunk]:
    """执行数据获取、清洗、分块、向量化，并可写入 Chroma。"""
    path = validate_document_path(file_path)
    text = read_docx_text(path)
    chunks = split_text(
        text,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=separators,
    )
    text_chunks = build_text_chunks(chunks, path)
    vectorized_chunks = vectorize_chunks(
        text_chunks,
        embeddings=embeddings,
        model_name=model_name,
    )
    if store_to_chroma:
        store_chunks_to_chroma(
            vectorized_chunks,
            persist_directory=persist_directory,
            collection_name=collection_name,
            clear_existing=clear_existing,
        )
    return vectorized_chunks
