from __future__ import annotations

"""文档读取对比 Demo

使用两种方式读取 `docx` 文档，并将结果输出到控制台，便于对比：
1. `python-docx`
2. `langchain-community` 的 `UnstructuredWordDocumentLoader`

说明：
- 这里主要用于对比 `docx` 读取效果
- `.doc` 文件对比通常更依赖 `UnstructuredWordDocumentLoader`
"""

from pathlib import Path

from langchain_core.documents import Document

# 直接修改这里来控制运行路径
INPUT_PATH = r"F:\PythonFiles\customer_service\测试文档.docx"

SUPPORTED_SUFFIXES = {".doc", ".docx"}


def collect_input_files(input_path: str) -> list[Path]:
    path = Path(input_path).expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(f"输入路径不存在: {path}")

    if path.is_file():
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            raise ValueError(f"仅支持 doc/docx 文件: {path}")
        return [path]

    files = [p for p in path.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES]
    if not files:
        raise ValueError(f"目录中未找到 doc/docx 文件: {path}")
    return sorted(files)


def read_with_python_docx(file_path: Path) -> str:
    """使用 python-docx 读取文档纯文本。"""
    from docx import Document as DocxDocument  # type: ignore[import-not-found]

    docx = DocxDocument(str(file_path))
    paragraphs = [p.text.strip() for p in docx.paragraphs if p.text.strip()]
    return "\n".join(paragraphs).strip()


def read_with_unstructured(file_path: Path) -> str:
    """使用 langchain-community 的 UnstructuredWordDocumentLoader 读取文档纯文本。"""
    try:
        from langchain_community.document_loaders import UnstructuredWordDocumentLoader
    except ImportError as exc:  # pragma: no cover
        raise ImportError(
            "缺少依赖：langchain-community。请先安装 `langchain-community` 和相关文档解析依赖。"
        ) from exc

    loader = UnstructuredWordDocumentLoader(str(file_path), mode="single")
    documents = loader.load()
    return "\n".join(doc.page_content.strip() for doc in documents if doc.page_content.strip()).strip()


def wrap_text(source: str, text: str, file_path: Path) -> Document:
    return Document(page_content=text, metadata={"source": source, "file": str(file_path)})


def print_result(title: str, document: Document) -> None:
    print(f"\n===== {title} =====")
    print(f"文件: {document.metadata.get('file', 'unknown')}")
    print(document.page_content if document.page_content else "[空内容]")


def main() -> None:
    files = collect_input_files(INPUT_PATH)

    for file_path in files:
        print(f"\n\n########################################")
        print(f"正在处理: {file_path}")
        print("########################################")

        if file_path.suffix.lower() != ".docx":
            print("[提示] python-docx 主要适用于 .docx，对 .doc 文件通常不适用。")

        try:
            python_docx_text = read_with_python_docx(file_path)
            print_result("python-docx 读取结果", wrap_text("python-docx", python_docx_text, file_path))
        except Exception as exc:
            print(f"\n===== python-docx 读取结果 =====\n文件: {file_path}\n[读取失败] {exc}")

        try:
            unstructured_text = read_with_unstructured(file_path)
            print_result(
                "UnstructuredWordDocumentLoader 读取结果",
                wrap_text("UnstructuredWordDocumentLoader", unstructured_text, file_path),
            )
        except Exception as exc:
            print(
                f"\n===== UnstructuredWordDocumentLoader 读取结果 =====\n文件: {file_path}\n[读取失败] {exc}"
            )


if __name__ == "__main__":
    main()
